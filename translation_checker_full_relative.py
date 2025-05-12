
import re
import os
import pandas as pd
from docx import Document

def extract_id_and_text_regex(paragraphs):
    extracted = []
    pattern = re.compile(r'\[(ID_\d+)\]\s*(.*)')
    for p in paragraphs:
        match = pattern.match(p)
        if match:
            extracted.append({'ID': match.group(1), 'Text': match.group(2)})
    return extracted

def split_combined_word_doc(input_docx_path):
    doc = Document(input_docx_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    korean_paragraphs, english_paragraphs = [], []
    id_counter = 1
    i = 0
    while i < len(paragraphs) - 1:
        korean_text = paragraphs[i]
        english_text = paragraphs[i + 1]
        korean_paragraphs.append(f'[ID_{id_counter}] {korean_text}')
        english_paragraphs.append(f'[ID_{id_counter}] {english_text}')
        id_counter += 1
        i += 2

    korean_doc = Document()
    english_doc = Document()
    for kp in korean_paragraphs:
        korean_doc.add_paragraph(kp)
    for ep in english_paragraphs:
        english_doc.add_paragraph(ep)

    korean_output_path = os.path.join(os.path.dirname(input_docx_path), 'korean_separated.docx')
    english_output_path = os.path.join(os.path.dirname(input_docx_path), 'english_separated.docx')
    korean_doc.save(korean_output_path)
    english_doc.save(english_output_path)

    return korean_output_path, english_output_path

def create_merged_excel_from_word(korean_path, english_path, output_excel_path):
    korean_doc = Document(korean_path)
    english_doc = Document(english_path)

    korean_paragraphs = [p.text.strip() for p in korean_doc.paragraphs if p.text.strip()]
    english_paragraphs = [p.text.strip() for p in english_doc.paragraphs if p.text.strip()]

    korean_data = extract_id_and_text_regex(korean_paragraphs)
    english_data = extract_id_and_text_regex(english_paragraphs)

    korean_df = pd.DataFrame(korean_data)
    english_df = pd.DataFrame(english_data)

    merged_df = pd.merge(english_df, korean_df, on='ID', suffixes=('_English', '_Korean'))
    merged_df.to_excel(output_excel_path, index=False)

    return output_excel_path, merged_df

def generate_translation_check_report(merged_df, glossary_path, output_excel_path):
    glossary_df = pd.read_csv(glossary_path)
    glossary_df = glossary_df.dropna(subset=['english', 'korean'])

    merged_df['Missing_English_Terms'] = ''
    merged_df['Expected_Korean_Terms'] = ''

    for index, row in merged_df.iterrows():
        english_text = row['Text_English']
        korean_text = row['Text_Korean']
        missing_terms = []
        expected_translations = []
        for _, glossary_row in glossary_df.iterrows():
            eng_term = glossary_row['english']
            kor_term = glossary_row['korean']
            if eng_term in english_text:
                if kor_term not in korean_text:
                    missing_terms.append(eng_term)
                    expected_translations.append(kor_term)
        merged_df.at[index, 'Missing_English_Terms'] = ', '.join(missing_terms) if missing_terms else 'None'
        merged_df.at[index, 'Expected_Korean_Terms'] = ', '.join(expected_translations) if expected_translations else 'None'

    expanded_rows = []
    for _, row in merged_df.iterrows():
        missing_terms = row['Missing_English_Terms'].split(', ') if row['Missing_English_Terms'] != 'None' else []
        expected_terms = row['Expected_Korean_Terms'].split(', ') if row['Expected_Korean_Terms'] != 'None' else []
        max_len = max(len(missing_terms), len(expected_terms), 1)
        for i in range(max_len):
            if missing_terms or expected_terms:
                expanded_rows.append({
                    'ID': row['ID'],
                    'Text_English': row['Text_English'],
                    'Text_Korean': row['Text_Korean'],
                    'Missing_English_Term': missing_terms[i] if i < len(missing_terms) else '',
                    'Expected_Korean_Term': expected_terms[i] if i < len(expected_terms) else ''
                })

    report_df = pd.DataFrame(expanded_rows)
    report_df.to_excel(output_excel_path, index=False)

    return output_excel_path

if __name__ == "__main__":
    base_dir = os.path.dirname(os.path.abspath(__file__))

    input_docx_path = os.path.join(base_dir, "input_combined.docx")
    glossary_path = os.path.join(base_dir, "glossary.csv")
    output_excel_path = os.path.join(base_dir, "merged_output.xlsx")
    report_output_path = os.path.join(base_dir, "translation_check_report.xlsx")

    korean_path, english_path = split_combined_word_doc(input_docx_path)
    merged_excel_path, merged_df = create_merged_excel_from_word(korean_path, english_path, output_excel_path)
    report_path = generate_translation_check_report(merged_df, glossary_path, report_output_path)

    print("✅ 변환 및 번역 체크 완료!")
    print(f" - 병합 엑셀: {merged_excel_path}")
    print(f" - 체크 리포트: {report_path}")
